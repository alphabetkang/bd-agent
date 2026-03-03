"""Process user-provided data sources into vector store documents."""
import io
import logging
from datetime import datetime

import feedparser
import httpx
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from ingestion.vector_store import get_vector_store

logger = logging.getLogger(__name__)

_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=80)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _fetch(url: str) -> bytes:
    with httpx.Client(follow_redirects=True, timeout=15) as client:
        resp = client.get(
            url, headers={"User-Agent": "Mozilla/5.0 (compatible; bd-agent/1.0)"}
        )
        resp.raise_for_status()
        return resp.content


def _ingest_docs(docs: list[Document], source_id: str) -> int:
    """Tag docs with source_id and upsert into vector store."""
    for doc in docs:
        doc.metadata["source_id"] = source_id
    store = get_vector_store()
    ids = [f"{source_id}_{i}" for i in range(len(docs))]
    store.add_documents(docs, ids=ids)
    return len(docs)


# ---------------------------------------------------------------------------
# RSS feed
# ---------------------------------------------------------------------------

def _process_rss(raw: bytes, source_name: str, url: str) -> list[Document]:
    feed = feedparser.parse(raw)
    docs = []
    for entry in feed.entries:
        title = entry.get("title", "")
        summary = entry.get("summary", "")
        content_list = entry.get("content", [])
        body = content_list[0].get("value", "") if content_list else summary
        text = f"{title}\n\n{body}".strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": source_name,
                    "url": entry.get("link", url),
                    "title": title,
                    "published": entry.get("published", ""),
                    "ingested_at": datetime.utcnow().isoformat(),
                },
            )
        )
    return docs


# ---------------------------------------------------------------------------
# Web page
# ---------------------------------------------------------------------------

def _process_webpage(raw: bytes, source_name: str, url: str) -> list[Document]:
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    title = soup.title.string.strip() if soup.title and soup.title.string else source_name
    text = soup.get_text(separator="\n", strip=True)
    if not text:
        return []
    chunks = _splitter.create_documents(
        [text],
        metadatas=[
            {
                "source": source_name,
                "url": url,
                "title": title,
                "ingested_at": datetime.utcnow().isoformat(),
            }
        ],
    )
    return chunks


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _process_pdf(content: bytes, filename: str) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": filename,
                    "page": i + 1,
                    "title": filename,
                    "ingested_at": datetime.utcnow().isoformat(),
                },
            )
        )
    return _splitter.split_documents(docs)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _process_docx(content: bytes, filename: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text:
        return []
    return _splitter.create_documents(
        [text],
        metadatas=[
            {
                "source": filename,
                "title": filename,
                "ingested_at": datetime.utcnow().isoformat(),
            }
        ],
    )


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def process_url(url: str, source_id: str, source_name: str) -> int:
    """Fetch a URL, detect type (RSS vs web page), and ingest into vector store."""
    raw = _fetch(url)
    feed = feedparser.parse(raw)
    if feed.entries:
        docs = _process_rss(raw, source_name, url)
        logger.info("URL treated as RSS: %d entries", len(docs))
    else:
        docs = _process_webpage(raw, source_name, url)
        logger.info("URL treated as web page: %d chunks", len(docs))

    if not docs:
        raise ValueError("No content could be extracted from the URL.")
    return _ingest_docs(docs, source_id)


def process_file(filename: str, content: bytes, source_id: str) -> int:
    """Process an uploaded file (PDF or DOCX) and ingest into vector store."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        docs = _process_pdf(content, filename)
    elif lower.endswith(".docx"):
        docs = _process_docx(content, filename)
    elif lower.endswith(".doc"):
        raise ValueError(".doc files are not supported. Please convert to .docx or .pdf.")
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    if not docs:
        raise ValueError("No text could be extracted from the file.")
    return _ingest_docs(docs, source_id)


def delete_source(source_id: str) -> None:
    """Remove all vector store documents belonging to a source."""
    store = get_vector_store()
    collection = store._collection
    results = collection.get(where={"source_id": source_id})
    if results and results.get("ids"):
        collection.delete(ids=results["ids"])
        logger.info("Deleted %d chunks for source %s", len(results["ids"]), source_id)
